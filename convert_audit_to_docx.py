"""
Convert paper_audit.txt → paper_audit.docx
Clean, readable reference document for programmers.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

INPUT  = "paper_audit.txt"
OUTPUT = "paper_audit.docx"

# ── Colour palette ──────────────────────────────────────────────────────────
C_DOMAIN_BG   = RGBColor(0x1a, 0x1a, 0x2e)   # dark navy – domain header bg
C_DOMAIN_FG   = RGBColor(0xFF, 0xFF, 0xFF)   # white text on navy
C_PAPER_BG    = RGBColor(0xEC, 0xF0, 0xF1)   # light grey – paper header bg
C_PAPER_FG    = RGBColor(0x1a, 0x1a, 0x2e)   # dark text on grey
C_LABEL       = RGBColor(0x27, 0x6F, 0xBF)   # blue – dimension labels
C_CORRECT_YES = RGBColor(0x1E, 0x8B, 0x4E)   # green
C_CORRECT_PAR = RGBColor(0xE6, 0x7E, 0x22)   # orange
C_CORRECT_NO  = RGBColor(0xC0, 0x39, 0x2B)   # red
C_CORRECT_NA  = RGBColor(0x7F, 0x8C, 0x8D)   # grey
C_BETTER_BG   = RGBColor(0xFF, 0xF9, 0xE6)   # pale yellow – Better Use bg
C_REC_BG      = RGBColor(0xE8, 0xF4, 0xFD)   # pale blue – Recommendation bg
C_TITLE_FG    = RGBColor(0x1a, 0x1a, 0x2e)
C_SUMMARY_BG  = RGBColor(0xF0, 0xF0, 0xF0)


def set_cell_bg(cell, color: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  rgb_hex(color))
    tcPr.append(shd)


def rgb_hex(color: RGBColor) -> str:
    """Convert RGBColor to 6-char hex string."""
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def set_para_bg(para, color: RGBColor):
    """Apply shading to a paragraph via pPr > shd."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  rgb_hex(color))
    pPr.append(shd)


def add_bookmark_anchor(para, name: str):
    """Add a named bookmark to a paragraph (for TOC links)."""
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"),   name[-8:].zfill(8))
    bm_start.set(qn("w:name"), name)
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), name[-8:].zfill(8))
    para._p.append(bm_start)
    para._p.append(bm_end)


def correct_color(text: str) -> RGBColor:
    t = text.strip().upper()
    if t.startswith("YES"):   return C_CORRECT_YES
    if t.startswith("PARTIAL"): return C_CORRECT_PAR
    if t.startswith("NO"):    return C_CORRECT_NO
    return C_CORRECT_NA   # N/A or unknown


# ── Document setup ───────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default body style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)


# ── Helper: add a paragraph with inline bold labels ────────────────────────
def add_body(text: str, indent: bool = False, bg: RGBColor = None):
    """Add a wrapped paragraph. Supports **bold** and [[correct:...]] markers."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after  = Pt(3)
    para.paragraph_format.space_before = Pt(0)
    if indent:
        para.paragraph_format.left_indent = Cm(0.6)
    if bg:
        set_para_bg(para, bg)

    # Simple text — no special markers
    run = para.add_run(text)
    run.font.size = Pt(10.5)
    return para


def add_dimension(label: str, content: str, bg: RGBColor = None):
    """One paragraph: [LABEL] content. Label in blue bold."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after  = Pt(4)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.left_indent  = Cm(0.4)
    if bg:
        set_para_bg(para, bg)

    r_label = para.add_run(f"{label}: ")
    r_label.bold = True
    r_label.font.color.rgb = C_LABEL
    r_label.font.size = Pt(10.5)

    r_text = para.add_run(content.strip())
    r_text.font.size = Pt(10.5)
    return para


def add_correct_dim(label: str, content: str):
    """Correct? line — coloured badge + content."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after  = Pt(4)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.left_indent  = Cm(0.4)

    r_label = para.add_run(f"{label}: ")
    r_label.bold = True
    r_label.font.color.rgb = C_LABEL
    r_label.font.size = Pt(10.5)

    # badge word (YES / PARTIAL / NO / N/A)
    first_word = content.strip().split(".")[0].split()[0] if content.strip() else ""
    badge_color = correct_color(first_word)

    r_badge = para.add_run(first_word + " ")
    r_badge.bold = True
    r_badge.font.color.rgb = badge_color
    r_badge.font.size = Pt(10.5)

    rest = content.strip()[len(first_word):].strip()
    if rest:
        r_rest = para.add_run(rest)
        r_rest.font.size = Pt(10.5)
    return para


def add_better_use(content: str):
    """Better Use: indented pale-yellow shaded block."""
    # Shaded label row
    p_label = doc.add_paragraph()
    p_label.paragraph_format.space_before = Pt(2)
    p_label.paragraph_format.space_after  = Pt(0)
    p_label.paragraph_format.left_indent  = Cm(0.4)
    set_para_bg(p_label, C_BETTER_BG)
    rl = p_label.add_run("Better Use: ")
    rl.bold = True
    rl.font.color.rgb = RGBColor(0xA0, 0x5C, 0x00)
    rl.font.size = Pt(10.5)

    # Content on same background
    lines = content.strip().split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Cm(1.0)
        set_para_bg(p, C_BETTER_BG)
        r = p.add_run(line)
        r.font.size = Pt(10.5)


def domain_header(text: str):
    """Dark navy full-width header for a domain section."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after  = Pt(6)
    set_para_bg(para, C_DOMAIN_BG)
    run = para.add_run(text)
    run.bold = True
    run.font.color.rgb = C_DOMAIN_FG
    run.font.size = Pt(13)
    run.font.name = "Calibri"


def paper_header(ref: str, citation: str):
    """Light grey bar for each paper: [P##] Author — Title."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(4)
    set_para_bg(para, C_PAPER_BG)

    r_ref = para.add_run(f"{ref}  ")
    r_ref.bold = True
    r_ref.font.color.rgb = C_PAPER_FG
    r_ref.font.size = Pt(11)

    r_cite = para.add_run(citation)
    r_cite.font.color.rgb = C_PAPER_FG
    r_cite.font.size = Pt(11)


def add_separator():
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "CCCCCC")
    pBdr.append(bot)
    pPr.append(pBdr)


def recommendation_block(ref: str, citation: str, content: str):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(2)
    set_para_bg(para, C_REC_BG)
    r1 = para.add_run(f"{ref}  ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x15, 0x4C, 0x79)
    r1.font.size = Pt(11)
    r2 = para.add_run(citation)
    r2.font.color.rgb = RGBColor(0x15, 0x4C, 0x79)
    r2.font.size = Pt(11)

    for line in content.strip().split("\n"):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Cm(0.6)
        set_para_bg(p, C_REC_BG)
        r = p.add_run(line)
        r.font.size = Pt(10.5)


def summary_block(text: str):
    for line in text.strip().split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        # Bold ALL-CAPS headings
        if line.isupper() or line.endswith(":") and line == line.upper().replace(":", ":"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            set_para_bg(p, C_SUMMARY_BG)
            r = p.add_run(line)
            r.bold = True
            r.font.size = Pt(11)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Cm(0.4)
            r = p.add_run(line)
            r.font.size = Pt(10.5)


# ── Parse the text file ─────────────────────────────────────────────────────
with open(INPUT, encoding="utf-8") as f:
    raw = f.read()

lines = raw.split("\n")

# ── Title page ───────────────────────────────────────────────────────────────
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(30)
p_title.paragraph_format.space_after  = Pt(6)
rt = p_title.add_run("MORRIGAN RESEARCH BIBLE")
rt.bold = True
rt.font.size = Pt(22)
rt.font.color.rgb = C_TITLE_FG

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(4)
rs = p_sub.add_run("Full Paper Audit — 7-Dimension Analysis")
rs.font.size = Pt(13)
rs.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p_key = doc.add_paragraph()
p_key.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_key.paragraph_format.space_after = Pt(4)
rk = p_key.add_run("Correct? key:  ")
rk.font.size = Pt(11)

for label, color in [("YES ", C_CORRECT_YES), ("PARTIAL ", C_CORRECT_PAR), ("NO ", C_CORRECT_NO), ("N/A", C_CORRECT_NA)]:
    rr = p_key.add_run(label + "  ")
    rr.bold = True
    rr.font.color.rgb = color
    rr.font.size = Pt(11)

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_date.paragraph_format.space_after = Pt(20)
rd = p_date.add_run("Generated: 2026-03-05   |   Papers: 81 cited (P1–P102)   |   90 audit entries")
rd.font.size = Pt(10)
rd.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

add_separator()

# ── State machine parser ─────────────────────────────────────────────────────
DIMS = ["Actual Insight", "Intended Use", "Domain/Context",
        "Examples", "Our Use", "Correct?", "Better Use"]

# Regex patterns
RE_DOMAIN   = re.compile(r"^={10,}$")
RE_DOMAIN_T = re.compile(r"^(DOMAIN \d+[^=]+|NEW PAPER RECOMMENDATIONS|AUDIT SUMMARY)$")
RE_PAPER    = re.compile(r"^\[P(\d+[a-z]?)\]\s+(.+)$")
RE_REC      = re.compile(r"^\[(REC-\d+)\]\s+(.+)$")
RE_DIM      = re.compile(r"^(Actual Insight|Intended Use|Domain/Context|Examples|Our Use|Correct\?|Better Use):\s*(.*)")
RE_SEP      = re.compile(r"^-{20,}$")

i = 0
n = len(lines)

in_summary        = False
in_rec_section    = False
current_dim       = None
current_dim_buf   = []
current_paper     = False

# buffer for multi-line dimensions
def flush_dim():
    global current_dim, current_dim_buf
    if current_dim is None:
        return
    content = " ".join(current_dim_buf).strip()
    if current_dim == "Better Use":
        # Split on numbered items: (1) ... (2) ...
        # Reformat as separate lines
        expanded = re.sub(r'\s*\((\d+)\)\s+', lambda m: f"\n({m.group(1)}) ", content).strip()
        add_better_use(expanded)
    elif current_dim == "Correct?":
        add_correct_dim(current_dim, content)
    else:
        add_dimension(current_dim, content)
    current_dim     = None
    current_dim_buf = []


skipping_equals = False
domain_text_pending = None

while i < n:
    line = lines[i]

    # Skip pure === separator lines — but capture following domain text
    if RE_DOMAIN.match(line.strip()):
        i += 1
        skipping_equals = True
        continue

    # If we just passed a === line, next non-empty line is probably domain title
    if skipping_equals:
        stripped = line.strip()
        if stripped == "":
            i += 1
            continue
        skipping_equals = False
        # Is it a domain title?
        if RE_DOMAIN_T.match(stripped):
            flush_dim()
            current_paper = False
            if "AUDIT SUMMARY" in stripped:
                in_summary = True
                # Heading
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(16)
                p.paragraph_format.space_after  = Pt(8)
                r = p.add_run("AUDIT SUMMARY")
                r.bold = True
                r.font.size = Pt(16)
                r.font.color.rgb = C_TITLE_FG
            elif "NEW PAPER RECOMMENDATIONS" in stripped:
                in_rec_section = True
                domain_header("NEW PAPER RECOMMENDATIONS")
            else:
                domain_header(stripped)
            i += 1
            continue
        # Otherwise process normally below
        # fall through with current line

    stripped = line.strip()

    # Blank line — separator between items
    if stripped == "" or RE_SEP.match(stripped):
        i += 1
        continue

    # Summary section — just dump
    if in_summary:
        flush_dim()
        summary_block(stripped)
        i += 1
        continue

    # Paper header [P##]
    m_paper = RE_PAPER.match(stripped)
    if m_paper:
        flush_dim()
        current_paper = True
        ref      = f"[P{m_paper.group(1)}]"
        citation = m_paper.group(2)
        paper_header(ref, citation)
        i += 1
        continue

    # Recommendation [REC-#]
    m_rec = RE_REC.match(stripped)
    if m_rec:
        flush_dim()
        current_paper = False
        # Collect multi-line body
        rec_ref  = f"[{m_rec.group(1)}]"
        rec_cite = m_rec.group(2)
        i += 1
        body_lines = []
        while i < n:
            nxt = lines[i].strip()
            # Stop at next recommendation, domain header, or separator
            if RE_REC.match(nxt) or RE_DOMAIN.match(nxt) or RE_DOMAIN_T.match(nxt):
                break
            body_lines.append(nxt)
            i += 1
        recommendation_block(rec_ref, rec_cite, "\n".join(body_lines))
        continue

    # Dimension line
    m_dim = RE_DIM.match(stripped)
    if m_dim:
        flush_dim()
        current_dim     = m_dim.group(1)
        first_content   = m_dim.group(2).strip()
        current_dim_buf = [first_content] if first_content else []
        i += 1
        continue

    # Continuation of current dimension
    if current_dim:
        current_dim_buf.append(stripped)
        i += 1
        continue

    # Anything else — plain paragraph
    add_body(stripped)
    i += 1

flush_dim()

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Saved -> {OUTPUT}")
